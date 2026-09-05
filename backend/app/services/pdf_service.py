from io import BytesIO
import os
import logging
from threading import Lock
from typing import Dict, List, Optional
import pymupdf
from ultralytics import YOLO
from huggingface_hub import hf_hub_download
# the below line must be imported before paddleocr package to resolve langchain import error
from app.patches.patch_langchain_imports import *
from app.services.docx_service import split_sentences, NUM_OF_SENTENCES_PER_SEGMENT
from paddleocr import PaddleOCR
from PIL import Image, ImageDraw
import numpy as np
import cv2
from itertools import groupby
from collections import Counter
import copy
import json
import tempfile
import re
from pdf2docx import Converter
import pdfplumber
from docx import Document as DocxDocument
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

_yolo_models: Dict[str, YOLO] = {}
_ocr_models: Dict[str, PaddleOCR] = {}
_model_init_lock = Lock()


def _nms_numpy(boxes: np.ndarray, scores: np.ndarray, iou_threshold: float = 0.5, ioa_threshold: float = 0.7) -> np.ndarray:
    """Apply non-maximum suppression and return kept indices.
    
    Suppresses boxes based on both IoU (symmetric overlap) and IoA
    (containment — catches cases where a small box sits mostly inside
    a larger one but IoU stays low due to size mismatch).
    """
    if boxes.size == 0:
        return np.array([], dtype=np.int64)

    x1 = boxes[:, 0]
    y1 = boxes[:, 1]
    x2 = boxes[:, 2]
    y2 = boxes[:, 3]
    areas = (x2 - x1) * (y2 - y1)
    order = scores.argsort()[::-1]

    keep = []
    while order.size > 0:
        i = int(order[0])
        keep.append(i)

        if order.size == 1:
            break

        rest = order[1:]

        xx1 = np.maximum(x1[i], x1[rest])
        yy1 = np.maximum(y1[i], y1[rest])
        xx2 = np.minimum(x2[i], x2[rest])
        yy2 = np.minimum(y2[i], y2[rest])

        w = np.maximum(0.0, xx2 - xx1)
        h = np.maximum(0.0, yy2 - yy1)
        inter = w * h

        # standard IoU
        union = areas[i] + areas[rest] - inter
        iou = np.divide(inter, union, out=np.zeros_like(inter), where=union > 0)

        # IoA: intersection over the SMALLER box's area
        min_area = np.minimum(areas[i], areas[rest])
        ioa = np.divide(inter, min_area, out=np.zeros_like(inter), where=min_area > 0)

        # suppress if either metric says "these are the same region"
        suppress = (iou > iou_threshold) | (ioa > ioa_threshold)
        order = rest[~suppress]

    return np.array(keep, dtype=np.int64)


def _get_yolo_model(device: str = "cpu") -> YOLO:
    """Lazy-load and cache YOLO model per device for process lifetime."""
    if device in _yolo_models:
        return _yolo_models[device]

    with _model_init_lock:
        if device in _yolo_models:
            return _yolo_models[device]

        filepath = hf_hub_download(
            repo_id="Armaggheddon/yolo26-document-layout",
            filename="yolo26s_doc_layout.pt",
            repo_type="model",
        )
        _yolo_models[device] = YOLO(filepath)
        return _yolo_models[device]


def _get_ocr_model(device: str = "cpu") -> PaddleOCR:
    """Lazy-load and cache PaddleOCR model per device for process lifetime."""
    if device in _ocr_models:
        return _ocr_models[device]

    with _model_init_lock:
        if device in _ocr_models:
            return _ocr_models[device]

        _ocr_models[device] = PaddleOCR(
            text_detection_model_name="PP-OCRv5_mobile_det",
            text_recognition_model_name="PP-OCRv5_mobile_rec",
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
            device=device,
        )
        return _ocr_models[device]

def pdf_to_images(pdf_bytes: bytes):
    """Generator that yields PIL Images one page at a time"""
    with pymupdf.open(stream=pdf_bytes, filetype="pdf") as doc:
        for page_num in range(len(doc)):
            page = doc[page_num]
            mat = pymupdf.Matrix(1, 1)
            pix = page.get_pixmap(matrix=mat)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            pix = None

            yield np.array(img)

def save_image_with_bbox(img_array: np.ndarray, bboxes, output_path: str, color = "red", width=3):
    """
    bbox: tuple (x1, y1, x2, y2)
    """
    img = Image.fromarray(img_array)
    draw = ImageDraw.Draw(img)
    colors = ["red", "blue", "green", "yellow", "purple", "orange", "pink", "cyan", "magenta", "lime", "teal", "lavender", "brown", "beige", "maroon",  "olive", "coral", "navy", "grey", "grey", "grey", "grey", "grey"]
    
    for i, box in enumerate(bboxes):
        draw.rectangle(box, outline=colors[i], width=width)
        # if i == 0:
        #     draw.rectangle(box, outline="red", width=width)
        # else:
        #     draw.rectangle(box, outline="blue", width=width)
        
        
    img.save(output_path)
    print(f"Saved to {output_path}")

def yolo_predict(image,c, device="cpu"):
    """
    Returns the result of doclayout-yolo
    The output is a tuple of:
        - a dict mapping classification labels with their description
        - tensor of shape (elements, 6) where elements is the number of detected elements
        and 6 is (x1 y1 x2 y2 conf classification)

    """
    model = _get_yolo_model(device)

    # Perform prediction
    result = model(
        image,
        conf=0.1,
        # disables multi-class classification
        agnostic_nms = True,
        device=device
    )

    # result[0].save(filename=f"result_{c}.jpg")
    # annotated_frame = result[0].plot(pil=True, line_width=5, font_size=20)
    # cv2.imwrite("result.jpg", annotated_frame)

    # before returning the data, remove overlapping boxes
    
    boxes = result[0].boxes.data[:, :4].detach().cpu().numpy()
    scores = result[0].boxes.conf.detach().cpu().numpy()
    classes = result[0].boxes.cls.detach().cpu().numpy()

    keep = _nms_numpy(boxes, scores, iou_threshold=0.5)

    filtered_boxes = boxes[keep]
    filtered_scores = scores[keep]
    filtered_classes = classes[keep]
    
    ## Testing
    # filtered_data = result[0].boxes.data[keep]
    # # 4. Overwrite the original object's data
    # result[0].boxes.data = filtered_data
    # # Annotate and save the result
    # result[0].save(filename=f"result_{c}.jpg")


    filtered = np.concatenate(
        (filtered_boxes, filtered_scores[:, np.newaxis], filtered_classes[:, np.newaxis]),
        axis=1,
    )
    
    
    # sort from up to bottom, left to right
    box_list = []
    for i in range(len(filtered)):
    # for i in range(len(result[0].boxes.data)): 
        bbox = filtered[i][:4].tolist()
        name_idx = filtered[i][-1].item()
        score = filtered[i][-2].item()
        box_list.append((bbox[0], bbox[1], bbox[2], bbox[3], score, name_idx))

    # Sort by y0 (top) then x0 (left) → reading order
    # box_list.sort(key=lambda x: (x[1], x[0]))

    return result[0].names, box_list


def ocr_predict(image, device="cpu"):
    """
    Arguments:
        - image, numpy.array
        
    Returns the result of paddleOCR
    
    The output is a tuple of:
        - List[str], a list of strings of extracted text, line by line
        - List[List[int]], 4 integers per text, representing (x1 y1 x2 y2)

        """
    ocr = _get_ocr_model(device)

    result = ocr.predict(image)

    return result[0]['rec_texts'], list([list(x) for x in result[0]["rec_boxes"]])

def ocr_in_yolo_ioa(ocr_box, yolo_box, threshold=0.5):
    """
    Check if either ocr_box is inside yolo box or yolo box is inside ocr_box, threshold is 0.5, meaning at least 50% of the smaller box should be inside the bigger box

    Returns:
        True or False
    """
    x1 = max(ocr_box[0], yolo_box[0])
    y1 = max(ocr_box[1], yolo_box[1])
    x2 = min(ocr_box[2], yolo_box[2])
    y2 = min(ocr_box[3], yolo_box[3])

    inter_w = max(0.0, x2 - x1)
    inter_h = max(0.0, y2 - y1)
    inter = inter_w * inter_h

    ocr_area = (ocr_box[2] - ocr_box[0]) * (ocr_box[3] - ocr_box[1])
    yolo_area = (yolo_box[2] - yolo_box[0]) * (yolo_box[3] - yolo_box[1])
    minimum_area = min(ocr_area, yolo_area)

    if inter / minimum_area >= threshold:
      return True
    
    return False


def extract_text_from_image(image, pdf_bytes, doc, c):
    """
    Extracts text from image

    Arguments:
        - image, numpy.array

    Returns:
        - ocr_text List[dict]: Each index represents a block of text in the page
        the dict has 2 keys:
            - 'text': represents the text of the block
            - 'bbox': represents the bounding box of the text, this is the bounding box of yolo with the text extracted using OCR
    """
    
    page = doc[c]
    ocr_result_texts, ocr_result_boxes= ocr_predict(image)
    yolo_result_names, yolo_result_data = yolo_predict(image, c)

    ocr_text = []
    tables = page.find_tables()
    break_flag = False
    # check if there are tables
    if tables.tables:
        for num, table in enumerate(tables.tables):
            data = table.extract()
            bbox = table.bbox
            rows = len(data)
            cols = max(len(r) for r in data)
            for i, row in enumerate(data):
                for j, cell in enumerate(row):
                    if cell.strip():
                        ocr_text.append(
                            {
                                "text": cell,
                                "bbox": bbox,
                                "type": "Table",
                                # goal of 'num' is to identify which table this cell belongs to
                                "info": {
                                    "num": num,
                                    "row": i,
                                    "col": j,
                                    "rows": rows,
                                    "cols": cols
                                }
                                
                            }
                        )

    for block_num in range(len(yolo_result_data)):
        # look for the classification label of this block, is it plain text or figure and so on
        name_idx = yolo_result_data[block_num][-1]
        
        block_bbox = yolo_result_data[block_num][:4]

        # we only need to extract text from those blocks
        # the only category that is out is picture
        # tables are extracted from pymupdf
        if yolo_result_names[name_idx] not in ['Caption',
                                            'Footnote',
                                            'Formula',
                                            'List-item',
                                            'Page-footer',
                                            'Section-header',
                                            'Text',
                                            'Title', 
                                            "Table"]:

            continue
        
        # if there are tables, check if the bounding boxes of those tables
        # overlap with the current block
        # if so, continue because i already extracted its text
        for block in ocr_text:
            if block["type"] == "Table":
                table_bbox = block["bbox"]
                if ocr_in_yolo_ioa(table_bbox, block_bbox, threshold=0.7):
                    break_flag = True  
                    break
                
                
        if break_flag:
            break_flag = False
            continue
        
        # try to get text
        block_text = page.get_text("text", clip=block_bbox)
        # image based content
        if not block_text.strip():
            for i in range(len(ocr_result_texts)):
                # check if small bounding box of the ocr is inside the bigger box of yolo
                if ocr_in_yolo_ioa(ocr_result_boxes[i], block_bbox):
                    block_text += ocr_result_texts[i] + "\n"


        # split sentences 
        sentences = split_sentences(block_text)
        for i in range(0, len(sentences), NUM_OF_SENTENCES_PER_SEGMENT):            
            group = " ".join(sentences[i:i + NUM_OF_SENTENCES_PER_SEGMENT])
            ocr_text.append(
            {
                "text": group,
                "bbox": block_bbox,
                # if yolo found a table, mark it as text because there is no table specific structure like the above pymupdf table extraction, so we will treat it as text
                "type": "Text" if yolo_result_names[name_idx] == "Table" else yolo_result_names[name_idx]
            }
            )

# the below code is for debugging purposes to visualize the extracted text and their bounding boxes, it can be removed later
    
    # Loop through OCR results and draw boxes
    # for item in ocr_text:
    #     text = item["text"]
    #     bbox = item["bbox"]

    #     # print(text, end="\n\n=============\n\n")
        
    #     # Convert bbox coordinates to integers
    #     # bbox format could be: [x1, y1, x2, y2] or [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
    #     if len(bbox) == 4 and all(isinstance(x, (int, float)) for x in bbox):
    #         # Rectangle format [x1, y1, x2, y2]
    #         x1, y1, x2, y2 = map(int, bbox)
    #         cv2.rectangle(image, (x1, y1), (x2, y2), (0, 255, 0), 2)
            
    #     elif len(bbox) == 4 and all(len(point) == 2 for point in bbox):
    #         # Polygon format [[x1,y1], [x2,y2], [x3,y3], [x4,y4]]
    #         pts = np.array(bbox, dtype=np.int32)
    #         cv2.polylines(image, [pts], True, (0, 255, 0), 2)
        
    
    # # Save or display result
    # cv2.imwrite(f"tmp/{c}.png", image)
    # cv2.imshow("Output", image)
    # cv2.waitKey(0)
    # cv2.destroyWindow("Output")
    ocr_text = sorted(ocr_text, key=lambda b: (round(b["bbox"][1], 0), b["bbox"][0]))
    return ocr_text

    
def extract_text_from_pdf(pdf_bytes: bytes):
    """
    Takes a pdf file and returns a List[List[dict]], outer index represent the different pages
    the inner index represent the different blocks of text inside a page

    Arguments:
        - pdf_file, bytes: a stream of bytes representing the pdf file

    Returns:
        - all_content, List[list[dict]]:
            The dict contains:
                - text, str: the text in the given block
                - bbox, tuple(int): x0, y0, x1, y1 -> the bounding boxes of the given text

    """
    all_content = []
    c = 0
    doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    for image in pdf_to_images(pdf_bytes):
        all_content.append(
            extract_text_from_image(image, pdf_bytes, doc, c)
        )
        c += 1
    
    return all_content

def _cell_is_empty(value) -> bool:
    """True if a pdfplumber cell has no real text content."""
    return value is None or str(value).strip() == ""


def _clean_pdfplumber_rows(rows: List[List[Optional[str]]]) -> List[List[Optional[str]]]:
    """
    Post-process a raw pdfplumber table.extract() result.

    - A row where every cell is empty/None is a blank grid-line artifact -> dropped.
    - A row with exactly one non-empty cell is a wrapped continuation line of a
      multi-line cell -> its text is appended (with a space) to the same-index
      cell of the row above, and the row itself is dropped.
    - Any other row (2+ non-empty cells) is a genuine row (header or data) and is
      kept as-is.
    """
    cleaned: List[List[Optional[str]]] = []

    for row in rows:
        non_empty_idx = [i for i, v in enumerate(row) if not _cell_is_empty(v)]

        if not non_empty_idx:
            # fully blank separator row
            continue

        if len(non_empty_idx) == 1 and cleaned:
            idx = non_empty_idx[0]
            text = str(row[idx]).strip()
            prev_row = cleaned[-1]
            if idx < len(prev_row):
                prev_val = prev_row[idx]
                if _cell_is_empty(prev_val):
                    prev_row[idx] = text
                else:
                    prev_row[idx] = f"{str(prev_val).rstrip()} {text}"
            else:
                # defensive: shouldn't normally happen, same table = same width
                continue
            continue

        cleaned.append(list(row))

    return cleaned


def _compact_table(cleaned_rows: List[List[Optional[str]]]) -> Optional[List[List[str]]]:
    """
    Drop the empty/spacer cells out of each cleaned row so that only the real
    column values remain (pdfplumber's grid detection produces several empty
    "filler" columns around each real column due to merged-cell borders).

    The resulting column count is taken as the most common compacted row
    length across the table; rows that disagree are padded/truncated to fit.
    Returns None if no consistent column count can be determined.
    """
    if not cleaned_rows:
        return None

    compacted_rows = [
        [str(v).strip() for v in row if not _cell_is_empty(v)] for row in cleaned_rows
    ]

    lengths = [len(r) for r in compacted_rows if r]
    if not lengths:
        return None
    target_len = Counter(lengths).most_common(1)[0][0]
    if target_len == 0:
        return None

    fitted_rows = []
    for row in compacted_rows:
        if len(row) < target_len:
            row = row + [""] * (target_len - len(row))
        elif len(row) > target_len:
            row = row[:target_len]
        fitted_rows.append(row)

    return fitted_rows


def _extract_pdfplumber_tables_by_page(pdf_bytes: bytes) -> List[List[List[List[str]]]]:
    """
    Returns, per page, a list of cleaned+compacted tables in document order.
    tables_by_page[page_index][table_index] -> List[List[str]] rows
    """
    tables_by_page: List[List[List[List[str]]]] = []
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages:
            page_tables = []
            for table in page.find_tables():
                raw_rows = table.extract()
                cleaned = _clean_pdfplumber_rows(raw_rows)
                compacted = _compact_table(cleaned)
                page_tables.append(compacted if compacted is not None else [])
            tables_by_page.append(page_tables)
    return tables_by_page


def _get_column_templates(table):
    """
    For each column, find a 'known good' paragraph template: a cell (any row)
    whose first paragraph already has a real run with explicit font
    properties. pdf2docx sometimes leaves a cell as an empty placeholder
    paragraph (no run, and a squashed `lineRule="exact"` 1pt line height)
    that originally just sat above a nested sub-table holding the real text.
    If we later drop text straight into that placeholder without fixing its
    paragraph properties, the text gets squashed into that ~1pt line box
    and rendered in Word's raw default font instead of the table's Arial.

    Returns a list (one entry per column) of (pPr_element_or_None,
    rPr_element_or_None) deep copies to use as a fallback for such cells.
    """
    num_cols = len(table.columns)
    templates = [None] * num_cols
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j >= num_cols or templates[j] is not None:
                continue
            para = cell.paragraphs[0]
            if para.runs and para.runs[0]._r.find(qn("w:rPr")) is not None:
                pPr = para._p.find(qn("w:pPr"))
                rPr = para.runs[0]._r.find(qn("w:rPr"))
                templates[j] = (
                    copy.deepcopy(pPr) if pPr is not None else None,
                    copy.deepcopy(rPr) if rPr is not None else None,
                )
    return templates


def _set_cell_text(cell, text: str, fallback_template=None) -> None:
    """
    Replace a docx table cell's text with `text`, reusing the formatting
    (font name/size/bold/italic) of the first existing run if there is one,
    so the cell keeps looking like the rest of the (pdf2docx-styled) table.

    pdf2docx sometimes represents a wrapped/stacked group of cells as a
    NESTED sub-table inside a single outer cell, rather than as separate
    outer rows (e.g. a "Product" header cell can contain a nested table
    holding "Verto Workspace"). cell.text/cell.paragraphs only see the
    outer paragraph text, so leaving those nested tables in place after we
    rewrite the outer text would leave stale, duplicated content behind.
    Since we're fully replacing this cell's content, any nested tables are
    removed too.

    If this cell's own paragraph has no run to copy formatting from (a sign
    it was previously an empty placeholder paragraph - see
    `_get_column_templates`), `fallback_template` - a (pPr, rPr) pair sourced
    from a known-good cell in the same column - is applied instead of
    leaving Word's raw defaults, which would otherwise squash the text into
    a ~1pt line box in the wrong font.
    """
    first_para = cell.paragraphs[0]
    template_run = first_para.runs[0] if first_para.runs else None

    # drop any nested sub-tables pdf2docx stashed inside this cell
    for nested_table in list(cell.tables):
        nested_table._tbl.getparent().remove(nested_table._tbl)

    # wipe all paragraphs except the first, and all runs within the first
    for para in cell.paragraphs[1:]:
        para._element.getparent().remove(para._element)
    for run in list(first_para.runs):
        run._element.getparent().remove(run._element)

    if template_run is None and fallback_template is not None:
        fallback_pPr, fallback_rPr = fallback_template
        existing_pPr = first_para._p.find(qn("w:pPr"))
        if existing_pPr is not None:
            first_para._p.remove(existing_pPr)
        if fallback_pPr is not None:
            first_para._p.insert(0, copy.deepcopy(fallback_pPr))

    new_run = first_para.add_run(text)
    if template_run is not None:
        new_run.font.name = template_run.font.name
        new_run.font.size = template_run.font.size
        new_run.font.bold = template_run.font.bold
        new_run.font.italic = template_run.font.italic
    elif fallback_template is not None and fallback_template[1] is not None:
        new_run._r.insert(0, copy.deepcopy(fallback_template[1]))


def _ensure_row_count(table, needed_rows: int) -> None:
    """Grow a docx table to `needed_rows` by cloning the last row's XML (and
    therefore its cell styling/borders/widths) as many times as needed."""
    last_tr = table.rows[-1]._tr
    while len(table.rows) < needed_rows:
        new_tr = copy.deepcopy(last_tr)
        table._tbl.append(new_tr)


def _rebuild_docx_table(table, new_rows: List[List[str]]) -> bool:
    """
    Overwrite `table`'s content with `new_rows`, reusing the table's existing
    style/borders/column widths. Returns False (no-op) if the column counts
    don't line up, since we can't safely "imitate the style" in that case.
    """
    if not new_rows:
        return False

    existing_cols = len(table.columns)
    new_cols = len(new_rows[0])
    if new_cols != existing_cols:
        return False

    # capture known-good per-column formatting BEFORE we start growing/
    # overwriting rows, so a placeholder cell's replacement text can borrow
    # sane formatting from a genuine cell elsewhere in the same column.
    column_templates = _get_column_templates(table)

    _ensure_row_count(table, len(new_rows))

    for i, row_data in enumerate(new_rows):
        for j, text in enumerate(row_data):
            _set_cell_text(table.rows[i].cells[j], text, fallback_template=column_templates[j])

    return True


def _replace_lossy_tables_with_pdfplumber(docx_bytes: bytes, pdf_bytes: bytes) -> bytes:
    """
    Compare every table pdf2docx produced against the equivalent table
    extracted (and cleaned) via pdfplumber, matched by page. Where
    pdfplumber recovered more rows (i.e. pdf2docx silently dropped/merged
    content), rebuild that table in the docx using the pdfplumber data.
    """
    tables_by_page = _extract_pdfplumber_tables_by_page(pdf_bytes)
    tables_per_page_counts = [len(page_tables) for page_tables in tables_by_page]

    document = DocxDocument(BytesIO(docx_bytes))
    docx_tables = document.tables

    if sum(tables_per_page_counts) != len(docx_tables):
        # Table counts between the two extraction methods don't line up, so we
        # can't safely match tables by page position. Bail out and keep the
        # original pdf2docx output untouched rather than risk touching the
        # wrong table.
        return docx_bytes

    docx_table_idx = 0
    for page_tables in tables_by_page:
        for pdfplumber_rows in page_tables:
            docx_table = docx_tables[docx_table_idx]
            docx_table_idx += 1

            docx_row_count = len(docx_table.rows)
            plumber_row_count = len(pdfplumber_rows)
            if plumber_row_count > docx_row_count:
                try:
                    _rebuild_docx_table(docx_table, pdfplumber_rows)
                except Exception:
                    # Never let a single bad table take down the whole
                    # conversion - just leave pdf2docx's version in place.
                    logger.exception(
                        "Failed to rebuild docx table %d from pdfplumber data; "
                        "keeping the original pdf2docx table instead.",
                        docx_table_idx - 1,
                    )
                    continue

    out = BytesIO()
    document.save(out)
    return out.getvalue()


def _convert_pdf_to_docx_bytes(pdf_bytes: bytes) -> bytes:
    """Convert PDF bytes to DOCX bytes using pdf2docx, then patch up any
    tables where pdf2docx dropped/merged rows by re-extracting them with
    pdfplumber."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_tmp:
        pdf_tmp.write(pdf_bytes)
        pdf_path = pdf_tmp.name

    docx_path = pdf_path.replace(".pdf", ".docx")
    try:
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()
    finally:
        os.unlink(pdf_path)
        if os.path.exists(docx_path):
            os.unlink(docx_path)

    try:
        docx_bytes = _replace_lossy_tables_with_pdfplumber(docx_bytes, pdf_bytes)
    except Exception:
        # Table touch-up is a best-effort improvement; never let it break the
        # base conversion.
        logger.exception(
            "pdfplumber table touch-up failed; returning the unmodified "
            "pdf2docx output instead."
        )

    return docx_bytes