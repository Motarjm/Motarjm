"""
This handles the reading of docx pages and generating them
"""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor
import re
from io import BytesIO

NUM_OF_SENTENCES_PER_SEGMENT = 1

def protect_false_periods(text: str) -> str:
    
    # Titles — always protect their period (always followed by a name)
    TITLES = {
        "mr", "mrs", "ms", "dr", "prof", "sr", "jr", "st",
        "sgt", "rev", "hon", "capt", "lt", "col", "gen",
    }

    # Non-title abbreviations — protect only when followed by lowercase/digit
    # (clearly mid-sentence). When followed by a capital they MAY end the
    # sentence, so we leave the period unprotected.
    SOFT_ABBREVS = {
        "etc", "e.g", "i.e", "vs", "approx", "dept", "est", "fig", "govt",
        "inc", "ltd", "no", "p", "pp", "vol",
        "jan", "feb", "mar", "apr", "jun", "jul", "aug", "sep", "oct", "nov", "dec",
    }

    # Known URL / file-extension TLDs
    URL_TLDS = {
        "com", "org", "net", "edu", "gov", "io", "co", "info", "biz",
        "pdf", "txt", "py", "js", "html", "htm", "csv", "json", "xml", "md",
    }
    
    # 1. Decimals and version numbers — e.g. 3.14, $4.99, 1,200.50, v1.2.3
    #    The multi-segment pattern catches 1.2.3 (which the old (\d)\.(\d)
    #    one-pass pattern missed because re.sub moves past the first match).
    text = re.sub(r'(\d+(?:\.\d+)+)',
                  lambda m: m.group(1).replace('.', '<<<DOT>>>'), text)

    # 1b. Email addresses — e.g. john.doe@example.com
    #     The internal dots between local-part labels must be protected,
    #     otherwise the no-whitespace splitter (pattern 2d) would split them.
    text = re.sub(r'\b([\w]+(?:\.[\w]+)*@[\w]+(?:\.[\w]+)+)\b',
                  lambda m: m.group(1).replace('.', '<<<DOT>>>'), text)

    # 2. URLs and file extensions — e.g. example.com, report.pdf
    tld_alt = '|'.join(re.escape(t) for t in URL_TLDS)
    text = re.sub(r'(\w)\.(' + tld_alt + r')\b', r'\1<<<DOT>>>\2', text)

    # 3. Title abbreviations — period always protected when followed by whitespace
    title_pat = r'\b(' + '|'.join(TITLES) + r')\.(?=\s)'
    text = re.sub(title_pat,
                  lambda m: m.group(0).replace('.', '<<<DOT>>>'),
                  text, flags=re.IGNORECASE)

    # 4a. Single-letter initial followed by another initial: "J. K."
    text = re.sub(r'\b([A-Z])\.(?=\s[A-Z]\.)', r'\1<<<DOT>>>', text)

    # 4b. Single-letter initial followed by a Capitalized word (a name):
    #     "J. Rowling" — but NOT when preceded by '.' or '>' (which would
    #     mean we're inside an already-protected acronym chain like U.S.).
    text = re.sub(r'(?<![.>])\b([A-Z])\.(?=\s[A-Z][a-z])', r'\1<<<DOT>>>', text)

    # 4c. Degree-style abbreviations: "Ph.D.", "M.D.", "B.A.", "M.S."
    #     Must run BEFORE step 4d so both dots are protected together.
    #     Protects BOTH dots so 'Ph.D. in physics' doesn't get split
    #     internally by the no-whitespace splitter (pattern 2d) later.
    text = re.sub(
        r'\b([A-Z][a-z]{1,3})\.([A-Z])\.(?=\s|$|[,;:!?)\]])',
        lambda m: m.group(0).replace('.', '<<<DOT>>>'),
        text,
    )

    # 4d. Single-letter initial followed by lowercase (clearly mid-sentence):
    #     "D. in physics" (the second half of Ph.D. after step 4c handles
    #     the pair), "U.S. in the world"
    text = re.sub(r'\b([A-Z])\.(?=\s[a-z])', r'\1<<<DOT>>>', text)

    # 5. Acronym chains — protect dots BETWEEN single capital letters
    #    (no space): "U.S.", "U.K.", "U.S.A." — but the trailing dot is
    #    left unprotected so it can still end a sentence.
    text = re.sub(r'\b([A-Z])\.(?=[A-Z])', r'\1<<<DOT>>>', text)

    # 6. Words ending in a capital letter when followed by lowercase.
    #    Covers ALL-CAPS words ("II. was", "NASA. was") AND mixed-case
    #    abbreviations ending in a capital ("PhD. gave", "McDonald. has").
    #    The heuristic: if the word's last letter is uppercase, it likely
    #    behaves like an abbreviation that doesn't end a sentence.
    text = re.sub(r'\b(\w*[A-Z])\.(?=\s[a-z])', r'\1<<<DOT>>>', text)

    # 7. Soft abbreviations — protect only when followed by lowercase/digit.
    #    IMPORTANT: use scoped (?i:...) so the abbreviation alternation is
    #    case-insensitive (matches Apr, APR, apr) BUT the lookahead
    #    (?=\s[a-z0-9]) stays case-sensitive — otherwise [a-z] would also
    #    match capitals, and 'Apr. We' would be wrongly protected.
    soft_pat = r'\b(?i:' + '|'.join(re.escape(a) for a in SOFT_ABBREVS) + r')\.(?=\s[a-z0-9])'
    text = re.sub(soft_pat,
                  lambda m: m.group(0).replace('.', '<<<DOT>>>'),
                  text)

    return text


def split_sentences(text: str) -> list[str]:
    # Step 1: Protect false periods
    text = protect_false_periods(text)

    # Step 2: Insert '\n' at every true sentence boundary

    # 2a. Multi-dot terminator (... or …) — split only when next char is a
    #     capital letter, digit, ¿/¡, accented capital, or CJK ideograph.
    text = re.sub(
        r'(\.{2,}|…+)(["\')\]\}]*)\s+'
        r'(?=[A-Z0-9¿¡\u00C0-\u00DE\u4e00-\u9fff])',
        lambda m: m.group(1) + m.group(2) + '\n',
        text,
    )

    # 2b. Single terminator (. ! ?) — not part of a multi-dot run — split
    #     on any non-space successor. Closing quotes/brackets/parens after
    #     the terminator are preserved on the prior sentence.
    text = re.sub(
        r'(?<![.\u2026])([.!?])(["\')\]\}]*)\s+(?=\S)',
        lambda m: m.group(1) + m.group(2) + '\n',
        text,
    )

    # 2c. CJK terminator (。！？) — split without requiring whitespace
    text = re.sub(
        r'([。！？])\s*(?=\S)',
        lambda m: m.group(1) + '\n',
        text,
    )

    # 2d. No-whitespace split: word-char + terminator + word-char
    #     Handles cases like "Hello.World" and "Wow!Amazing" where the
    #     writer forgot to put a space after the sentence-ending punctuation.
    text = re.sub(
        r'(?<=\w)([.!?])(?=\w)',
        lambda m: m.group(1) + '\n',
        text,
    )

    # Step 3: Split on newline and restore protected dots.
    # Filter out fragments that contain no alphanumeric character
    # (e.g. pure-punctuation fragments like '...' or '.').
    sentences = text.split('\n')
    sentences = [s.replace('<<<DOT>>>', '.').strip() for s in sentences]
    return [s for s in sentences if s and any(c.isalnum() for c in s)]


def iter_unique_cells_table(table):
    """Yield each cell in a table only once, handling both horizontal and vertical merges."""
    seen_tc = set()
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            tc = cell._tc
            if tc in seen_tc:
                continue
            seen_tc.add(tc)
            yield row_idx, col_idx, cell
            
def extract_run_format(paragraph) -> dict:
    """Grab the dominant run formatting for a paragraph (first non-empty run)."""
    for run in paragraph.runs:
        if run.text.strip():
            color = None
            try:
                if run.font.color and run.font.color.type is not None:
                    color = str(run.font.color.rgb) if run.font.color.rgb else None
            except Exception:
                color = None
            return {
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None,
                "bold": run.font.bold,
                "italic": run.font.italic,
                "underline": run.font.underline,
                "color": color,
            }
    return {}

            
def infer_block_type(paragraph) -> str:
    """Map a source paragraph's Word style to our layout-type vocabulary."""
    style_name = (paragraph.style.name or "").lower()
    if "title" in style_name:
        return "Title"
    if "heading" in style_name:
        return "Section-header"
    if "caption" in style_name:
        return "Caption"
    if "list" in style_name or paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
        return "List-item"
    return "Text" 


def resolve_effective_alignment(paragraph):
    """Direct paragraph formatting wins. Otherwise walk the style's
    base_style chain until an explicit alignment is found — Heading/Title
    styles almost never set alignment directly on themselves, they inherit
    it from a parent style, so checking only paragraph.style.paragraph_format
    (without walking base_style) misses this and silently skips the mirror
    for headings. 
    
    Returns None if no explicit alignment is found anywhere
    in the chain, which means Word is rendering it at its true default: left."""
    if paragraph.alignment is not None:
        return paragraph.alignment
    style = paragraph.style
    while style is not None:
        fmt = style.paragraph_format
        if fmt is not None and fmt.alignment is not None:
            return fmt.alignment
        style = style.base_style
    return None
 
def add_paragraph_blocks(blocks, paragraph, para_idx):
    """Same sentence-grouping as before, but each block keeps a reference
    to its source paragraph so translations can be written back into the
    ORIGINAL paragraph object — never a freshly built one."""
    text = paragraph.text.strip()
    if not text:
        return
    style_name = paragraph.style.name if paragraph.style else "Normal"
    run_fmt = extract_run_format(paragraph)
    btype = infer_block_type(paragraph)

    sentences = split_sentences(text)
    for i in range(0, len(sentences), NUM_OF_SENTENCES_PER_SEGMENT):
        group = " ".join(sentences[i:i + NUM_OF_SENTENCES_PER_SEGMENT])
        blocks.append({
            "id": f"{para_idx}_{i}",      # ← NEW
            "text": group,
            "type": btype,
            "bbox": [],
            "info": {"style_name": style_name, **run_fmt},
            "_paragraph_ref": paragraph,  # internal only — never sent to the frontend
        })


def add_table_blocks(blocks, table, table_num, table_idx):
    for row_idx, col_idx, cell in iter_unique_cells_table(table):
        text = cell.text.strip()
        if not text:
            continue
        cell_para = cell.paragraphs[0] if cell.paragraphs else None
        cell_fmt = extract_run_format(cell_para) if cell_para else {}

        sentences = split_sentences(text)
        for i in range(0, len(sentences), NUM_OF_SENTENCES_PER_SEGMENT):
            group = " ".join(sentences[i:i + NUM_OF_SENTENCES_PER_SEGMENT])
            blocks.append({
                "id": f"{table_idx}_{row_idx}_{col_idx}_{i}",   # add sentence index
                "text": group,
                "type": "Table",
                "bbox": [],
                "info": {"num": table_num, "row": row_idx, "col": col_idx, **cell_fmt},
                "_cell_ref": cell,
            })


def get_docx_blocks(docx_bytes: bytes):
    """Returns (doc, blocks) instead of just blocks — the caller must keep
    `doc` alive and pass it into apply_translations later. This is the same
    walk as before; the only change is we no longer discard the parsed
    Document once blocks are extracted."""
    doc = Document(docx_bytes)
    blocks = []
    table_num = 0

    def walk(container_element):
        nonlocal table_num
        for i, element in enumerate(container_element):
            if element.tag.endswith('p'):
                add_paragraph_blocks(blocks, Paragraph(element, doc), i)
            elif element.tag.endswith('tbl'):
                add_table_blocks(blocks, Table(element, doc), table_num, i)
                table_num += 1

    if doc.sections and doc.sections[0].header:
        walk(doc.sections[0].header._element)
    walk(doc.element.body)
    if doc.sections and doc.sections[0].footer:
        walk(doc.sections[0].footer._element)

    return doc, blocks


def apply_complex_script_font(run, font_name: str = "Arial"):
    """
    if the font lacks full Arabic coverage, unsupported glyphs render as
    the placeholder .notdef box, which looks exactly like words being
    replaced with dots. w:hint tells Word to actually prefer the cs slot
    for this run instead of guessing based on the first character."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), font_name)
    rFonts.set(qn('w:hint'), 'cs')


def _write_translation_into_paragraph(paragraph, translated_text, complex_font: str = "Arial"):
    """Overwrite text in place.
    
    Paragraph-level formatting (alignment,
    spacing, indentation) lives on the pPr element and is left untouched
    except for the bidi flag and alignment mirror applied below. 
    
    Run-level formatting (bold/italic/color/font) lives on run[0]'s rPr, which is
    also left untouched except for the RTL/complex-script fields added
    below — only run[0].text is replaced, so its formatting carries over
    unchanged. Extra runs are emptied rather than deleted, so no rPr/rIds
    are broken."""

    # ── 1. Paragraph-level RTL (w:pPr/w:bidi) ──
    pPr = paragraph._element.get_or_add_pPr()
    if pPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    # ── 2. Mirror alignment: left ↔ right ──
    # Uses resolve_effective_alignment instead of reading paragraph.alignment
    effective_align = resolve_effective_alignment(paragraph)
    if effective_align in (WD_ALIGN_PARAGRAPH.LEFT, None):
        # None means no explicit alignment anywhere in the chain, which is
        # Word's true rendered default (left) — treat it the same as LEFT.
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif effective_align == WD_ALIGN_PARAGRAPH.RIGHT:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # CENTER / JUSTIFY — leave as-is, no mirroring needed.

    # ── 3. Ensure at least one run exists ──
    runs = paragraph.runs
    if not runs:
        new_run = paragraph.add_run(translated_text)
        rPr = new_run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
        apply_complex_script_font(new_run, complex_font)
        return

    # ── 4. Overwrite existing runs ──
    runs[0].text = translated_text
    for run in runs[1:]:
        run.text = ""

    # ── 5. Run-level RTL (w:rPr/w:rtl) + complex-script font ──
    rPr = runs[0]._element.get_or_add_rPr()
    if rPr.find(qn('w:rtl')) is None:
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
    apply_complex_script_font(runs[0], complex_font)


def apply_translations(doc, blocks, complex_font: str = "Arial") -> bytes:
    """Writes every block's translated_text back into the ORIGINAL document.
    No Document(), no add_table, no resolve_style, no apply_run_format —
    table borders/shading/merges, images, headers, footers, and page setup
    are untouched because the document object they live on is never
    replaced, only edited."""
    para_groups, cell_groups = {}, {}
    for block in blocks:
        if "_paragraph_ref" in block:
            para_groups.setdefault(id(block["_paragraph_ref"]), []).append(block)
        elif "_cell_ref" in block:
            cell_groups.setdefault(id(block["_cell_ref"]), []).append(block)

    for group in para_groups.values():
        paragraph = group[0]["_paragraph_ref"]
        translated = " ".join(b.get("translated_text", b["text"]) for b in group)
        _write_translation_into_paragraph(paragraph, translated, complex_font)

    for group in cell_groups.values():
        cell = group[0]["_cell_ref"]
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        translated = " ".join(b.get("translated_text", b["text"]) for b in group)
        _write_translation_into_paragraph(paragraph, translated, complex_font)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def build_docx(original_docx_bytes: bytes, translated_contents: list[list[dict]]) -> bytes:
    """
    Re-parses the original docx and writes translations back by matching
    segment IDs. Multiple translated segments with the same ID (split case)
    are joined with spaces.
    """
    doc, blocks = get_docx_blocks(BytesIO(original_docx_bytes))

    # Build lookup by ID, joining splits with spaces
    translation_by_id = {}
    for t in translated_contents[0]:
        tid = t.get("id")
        if tid is not None:
            txt = t.get("translated_text", "")
            if tid in translation_by_id:
                translation_by_id[tid] += " " + txt
            else:
                translation_by_id[tid] = txt

    for block in blocks:
        tid = block.get("id")
        if tid is not None and tid in translation_by_id:
            block["translated_text"] = translation_by_id[tid]
        # If ID not found (e.g. merged away), apply_translations falls back to original text
    print(translation_by_id)
    return apply_translations(doc, blocks)

"""
=================================================
Below code is used if uploaded doc was PDF
"""
def apply_run_format(run, fmt: dict):
    if not fmt:
        return
    if fmt.get("font_name"):
        run.font.name = fmt["font_name"]
    if fmt.get("font_size"):
        run.font.size = Pt(fmt["font_size"])
    if fmt.get("bold") is not None:
        run.font.bold = fmt["bold"]
    if fmt.get("italic") is not None:
        run.font.italic = fmt["italic"]
    if fmt.get("underline") is not None:
        run.font.underline = fmt["underline"]
    if fmt.get("color"):
        try:
            run.font.color.rgb = RGBColor.from_string(fmt["color"])
        except Exception:
            pass

def resolve_style(doc, style_name, fallback=None):
    """Return style_name if it exists in this document's style set,
    otherwise fallback (or None, meaning 'use Word's default')."""
    if not style_name:
        return fallback
    try:
        doc.styles[style_name]
        return style_name
    except KeyError:
        return fallback
    
def add_table(doc, blocks):
    """
    Rebuilds a table from cell-level blocks, each carrying:
      block["info"] = {rows, cols, row, col, table_style, col_widths,
                        font_name, font_size, bold, italic, underline, color}
    """
    if not blocks:
        return
 
    info0 = blocks[0].get("info", {})
    n_rows = info0.get("rows", 0)
    n_cols = info0.get("cols", 0)
 
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style =  resolve_style(doc, info0.get("table_style"), "Table Grid")
 
    col_widths = info0.get("col_widths") or []
    for i, width in enumerate(col_widths):
        if width and i < len(table.columns):
            for cell in table.columns[i].cells:
                cell.width = width
 
    for block in blocks:
        info = block.get("info", {})
        row = info.get("row", 0)
        col = info.get("col", 0)
        val = str(block.get("translated_text", ""))
 
        cell = table.cell(row, col)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(val)
        apply_run_format(run, info)
        
        
def build_docx_from_scratch(pages):
    """
    pages: list (per page) of list (per block) of dicts like
           {"type": "Text", "text": "...", "bbox": [x0, y0, x1, y1], ...}
    """
    doc = Document()

    # --- Set up custom styles once ---
    styles = doc.styles
    
    def set_style_rtl(style):
        pPr = style.element.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    set_style_rtl(doc.styles['Normal'])
    footer_flag = False
    print(pages)
    for page_idx, blocks in enumerate(pages):
        footnotes_buffer = []
        page_table = []

        for block_num, block in enumerate(blocks):
            btype = block.get("type")
            
            
            text = block.get("translated_text", "").strip()
            text = text.replace("\n", " ").strip()  # Normalize newlines to spaces
            info = block.get("info") or {}
            
            if btype != "Table" and page_table:
                # Flush any accumulated table blocks before processing non-table blocks
                add_table(doc, page_table)
                page_table = []
                        
                        
            if not text:
                continue
            
            style_name = info.get("style_name")
                    


            if btype == "Title":
                doc.add_heading(text, level=0)

            elif btype == "Section-header":
                doc.add_heading(text, level=1)

            elif btype == "Text" or btype is None:
                p = doc.add_paragraph(text, style=resolve_style(doc, style_name))
                if not style_name:
                    pass
                for run in p.runs:
                    apply_run_format(run, info)

            elif btype == "List-item":
                p = doc.add_paragraph(text, style=resolve_style(doc, style_name, "List Bullet"))
                for run in p.runs:
                    apply_run_format(run, info)

            elif btype == "Caption":
                p = doc.add_paragraph(text, style=resolve_style(doc, style_name, "Caption"))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    apply_run_format(run, info)

            elif btype == "Formula":
                p = doc.add_paragraph(text, style= "Formula")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Cambria Math"

            elif btype == "Footnote":
                # Defer — collect and dump at bottom of the page's content
                footnotes_buffer.append(text)

            elif btype == "Page-footer":
                footer_flag = True
                footer = text
                
                
            elif btype == "Table":
                # if first table or another part of the same table
                if block.get("info", {}).get("num") == 0 \
                    or (block_num > 0 and
                        block.get("info", {}).get("num") == blocks[block_num - 1].get("info", {}).get("num")):
                    page_table.append(block)
                
                # new table
                else:
                    add_table(doc, page_table)
                    page_table = []
                    page_table.append(block)
                    
        # Flush any remaining table blocks at the end of the page
        if page_table:
            add_table(doc, page_table)
            page_table = []
         
         
        # Flush footnotes for this page as a small block at the end
        if footnotes_buffer:
            sep = doc.add_paragraph()
            sep.paragraph_format.space_before = Pt(12)
            for fn in footnotes_buffer:
                doc.add_paragraph(fn, style="Footnote")
                
        if page_idx < len(pages) - 1:
            doc.add_page_break()

    if footer_flag:
        doc.sections[0].footer.paragraphs[0].text = footer
        
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()